#!/usr/bin/env python3
"""
Create a Word document from a Donyati template using content specification.

Usage:
    python create_document.py content.json output.docx [--template standard|cover]

The content.json should have the structure:
{
  "title": "Document Title",
  "subtitle": "Optional subtitle",
  "client": "Client Name",
  "date": "January 2025",
  "confidential": false,
  "header_text": "Optional custom header text",
  "footer_text": "Optional custom footer text",
  "sections": [
    {
      "heading": "Section Title",
      "level": 1,
      "content": [
        {"type": "paragraph", "text": "Paragraph text..."},
        {"type": "bullets", "items": ["Item 1", "Item 2", "Item 3"]},
        {"type": "numbered", "items": ["Step 1", "Step 2", "Step 3"]},
        {"type": "table", "headers": ["Col1", "Col2"], "rows": [["A", "B"], ["C", "D"]]},
        {"type": "callout", "style": "info|warning|success", "title": "Note", "text": "Important info..."},
        {"type": "highlight_box", "text": "Key takeaway or metric", "style": "primary|accent"},
        {"type": "image", "path": "/path/to/image.png", "width": 6.5, "caption": "Optional caption"},
        {"type": "mermaid", "definition": "graph TD\\n    A-->B", "width": 6.5, "caption": "Optional caption"}
      ]
    }
  ]
}

Template Selection:
- Auto-selects cover page template if: confidential=true, or 4+ sections, or type is proposal/report
- Use --template flag to override: "standard" or "cover"
"""

import argparse
import json
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# Template paths (relative to DonyatiDocs root)
TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "word"
STANDARD_TEMPLATE = TEMPLATES_DIR / "donyati-word-template-2025.docx"
COVER_TEMPLATE = TEMPLATES_DIR / "donyati-word-template-cover-2025.docx"

# Donyati brand colors
DONYATI_BLACK = RGBColor(0x12, 0x00, 0x2A)
DONYATI_PURPLE = RGBColor(0x4A, 0x47, 0x78)
DONYATI_LIGHT_PURPLE = RGBColor(0x6B, 0x68, 0x99)

# Callout/box colors
COLORS = {
    'primary': {'bg': 'E8E6F0', 'border': '4A4778', 'text': '12002A'},
    'accent': {'bg': 'F5F0FA', 'border': '6B6899', 'text': '12002A'},
    'info': {'bg': 'E3F2FD', 'border': '1976D2', 'text': '0D47A1'},
    'warning': {'bg': 'FFF8E1', 'border': 'F9A825', 'text': '5D4037'},
    'success': {'bg': 'E8F5E9', 'border': '2E7D32', 'text': '1B5E20'},
}


def should_use_cover_template(content: Dict[str, Any]) -> bool:
    """Determine if cover page template should be used."""
    if content.get("confidential", False):
        return True
    doc_type = content.get("type", "").lower()
    if doc_type in ["proposal", "report", "deliverable", "assessment", "executive summary"]:
        return True
    sections = content.get("sections", [])
    if len(sections) >= 4:
        return True
    return False


def replace_text_in_element(element, old_text: str, new_text: str):
    """Recursively replace text in XML element and all children."""
    try:
        if hasattr(element, 'text') and element.text and old_text in element.text:
            try:
                element.text = element.text.replace(old_text, new_text)
            except AttributeError:
                pass  # Some elements don't allow text setting
        if hasattr(element, 'tail') and element.tail and old_text in element.tail:
            try:
                element.tail = element.tail.replace(old_text, new_text)
            except AttributeError:
                pass
    except TypeError:
        pass  # Skip elements that don't support text operations

    try:
        for child in element:
            replace_text_in_element(child, old_text, new_text)
    except TypeError:
        pass  # Element is not iterable


def replace_cover_page_placeholders(doc: Document, title: str, client: str):
    """Replace placeholders on the cover page."""
    # Replace in paragraphs and runs
    for para in doc.paragraphs:
        for run in para.runs:
            if "[Document Title]" in run.text:
                run.text = run.text.replace("[Document Title]", title)
            if "[Client Legal Name]" in run.text:
                run.text = run.text.replace("[Client Legal Name]", client if client else "")

    # Also search through the raw XML for text in shapes/textboxes
    # using XPath to find all text elements
    from lxml import etree
    nsmap_local = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
    }

    # Find all text elements in the document XML
    body_xml = doc._body._body
    for t_elem in body_xml.iter():
        tag_name = t_elem.tag.split('}')[-1] if '}' in str(t_elem.tag) else str(t_elem.tag)
        if tag_name == 't' and t_elem.text:
            if "[Document Title]" in t_elem.text:
                t_elem.text = t_elem.text.replace("[Document Title]", title)
            if "[Client Legal Name]" in t_elem.text:
                t_elem.text = t_elem.text.replace("[Client Legal Name]", client if client else "")


def update_header(doc: Document, header_text: Optional[str]):
    """Update header text by replacing placeholder text and setting Donyati purple color."""
    if not header_text:
        return

    # Both templates use different placeholder text
    placeholders = ["Enter Report Name Here", "Enter Document Name Here"]

    for section in doc.sections:
        header = section.header
        # Search in all XML text elements (including inside SDT content controls)
        found = False
        for elem in header._element.iter():
            if elem.tag.endswith('}t') and elem.text:
                for placeholder in placeholders:
                    if placeholder in elem.text:
                        elem.text = elem.text.replace(placeholder, header_text)
                        found = True

                        # Set color to Donyati purple on the parent run element
                        # Find the parent <w:r> (run) element and set color
                        parent = elem.getparent()
                        while parent is not None:
                            if parent.tag.endswith('}r'):
                                # Find or create rPr (run properties)
                                rPr = parent.find(qn('w:rPr'))
                                if rPr is None:
                                    rPr = OxmlElement('w:rPr')
                                    parent.insert(0, rPr)
                                # Find or create color element
                                color_elem = rPr.find(qn('w:color'))
                                if color_elem is None:
                                    color_elem = OxmlElement('w:color')
                                    rPr.append(color_elem)
                                # Set to Donyati purple
                                color_elem.set(qn('w:val'), '4A4778')
                                break
                            parent = parent.getparent()
                        break

        # If no placeholder found, add new paragraph
        if not found:
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(header_text)
            run.font.size = Pt(9)
            run.font.color.rgb = DONYATI_PURPLE


def update_footer(doc: Document, footer_text: Optional[str], date_text: Optional[str] = None):
    """Update footer text and date by replacing placeholders."""
    # The date placeholder is split: "Click or tap to enter a date" + "."
    date_placeholder = "Click or tap to enter a date"

    for section in doc.sections:
        footer = section.footer

        # Replace date placeholder if date provided
        # Search in all XML text elements (including inside SDT content controls)
        if date_text:
            found_date = False
            for elem in footer._element.iter():
                if elem.tag.endswith('}t') and elem.text:
                    if date_placeholder in elem.text:
                        elem.text = elem.text.replace(date_placeholder, date_text)
                        found_date = True
                    # Also remove the trailing period that's in a separate element
                    elif found_date and elem.text == '.':
                        elem.text = ''

        # Handle custom footer text (replaces center content if provided)
        if footer_text:
            # Look for the center paragraph (usually copyright) and update it
            found_center = False
            for para in footer.paragraphs:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER or "Donyati" in para.text or "©" in para.text:
                    # Clear and update
                    para.clear()
                    run = para.add_run(footer_text)
                    run.font.size = Pt(9)
                    run.font.color.rgb = DONYATI_PURPLE
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    found_center = True
                    break

            # If no center paragraph found, add one
            if not found_center and footer.paragraphs:
                para = footer.paragraphs[0]
                para.clear()
                run = para.add_run(footer_text)
                run.font.size = Pt(9)
                run.font.color.rgb = DONYATI_PURPLE
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with the appropriate style."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style = style_map.get(level, "Heading 1")
    doc.add_paragraph(text, style=style)


def add_paragraph(doc: Document, text: str):
    """Add a normal paragraph."""
    doc.add_paragraph(text, style="Normal")


def add_bullets(doc: Document, items: List[str]):
    """Add bullet points using Bullet 1 style."""
    for item in items:
        doc.add_paragraph(item, style="Bullet 1")


def add_numbered_list(doc: Document, items: List[str]):
    """Add numbered list."""
    for i, item in enumerate(items, 1):
        para = doc.add_paragraph(style="List Paragraph")
        para.add_run(f"{i}. {item}")


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a table with Donyati styling."""
    num_cols = len(headers)
    num_rows = len(rows) + 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Grid Table 4 Accent 6"

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        set_cell_shading(cell, COLORS['primary']['border'])
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        table_row = table.rows[row_idx + 1]
        bg_color = COLORS['primary']['bg'] if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table_row.cells[col_idx]
                cell.text = str(cell_text)
                set_cell_shading(cell, bg_color)

    doc.add_paragraph()


def add_callout_box(doc: Document, title: str, text: str, style: str = "info"):
    """Add a styled callout box (info, warning, success)."""
    colors = COLORS.get(style, COLORS['info'])

    # Create a single-cell table for the callout box
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False

    cell = table.rows[0].cells[0]

    # Set cell width to full page width
    cell.width = Inches(6.5)

    # Set cell shading
    set_cell_shading(cell, colors['bg'])

    # Set left border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '36')  # 4.5pt border
    left_border.set(qn('w:color'), colors['border'])
    tcBorders.append(left_border)

    # Set thin borders for other sides
    for side in ['top', 'right', 'bottom']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), colors['border'])
        tcBorders.append(border)

    tcPr.append(tcBorders)

    # Add title if provided
    if title:
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor.from_string(colors['text'])

        # Add text in new paragraph
        text_para = cell.add_paragraph()
        text_run = text_para.add_run(text)
        text_run.font.size = Pt(10)
        text_run.font.color.rgb = RGBColor.from_string(colors['text'])
    else:
        # Just add text
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(colors['text'])

    # Add spacing after
    doc.add_paragraph()


def add_highlight_box(doc: Document, text: str, style: str = "primary"):
    """Add a highlighted box for key metrics or takeaways."""
    colors = COLORS.get(style, COLORS['primary'])

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)

    # Set shading
    set_cell_shading(cell, colors['bg'])

    # Set all borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), colors['border'])
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # Add centered text
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(colors['text'])

    doc.add_paragraph()


def add_metric_row(doc: Document, metrics: List[Dict[str, str]]):
    """Add a row of metric boxes."""
    num_cols = len(metrics)
    table = doc.add_table(rows=1, cols=num_cols)
    table.autofit = False

    col_width = Inches(6.5 / num_cols)

    for i, metric in enumerate(metrics):
        cell = table.rows[0].cells[i]
        cell.width = col_width

        style = metric.get('style', 'primary')
        colors = COLORS.get(style, COLORS['primary'])
        set_cell_shading(cell, colors['bg'])

        # Set borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:color'), colors['border'])
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # Add value
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = para.add_run(metric.get('value', ''))
        value_run.bold = True
        value_run.font.size = Pt(24)
        value_run.font.color.rgb = RGBColor.from_string(colors['border'])

        # Add label
        label_para = cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.add_run(metric.get('label', ''))
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor.from_string(colors['text'])

    doc.add_paragraph()


def add_image(doc: Document, path: str, width: float = 6.5, caption: Optional[str] = None):
    """Add an image to the document with optional caption."""
    image_path = Path(path)
    if not image_path.exists():
        print(f"Warning: Image not found: {path}")
        add_paragraph(doc, f"[Image not found: {path}]")
        return

    doc.add_picture(str(image_path), width=Inches(width))

    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = DONYATI_PURPLE


def add_mermaid_diagram(doc: Document, definition: str, width: float = 6.5, caption: Optional[str] = None):
    """Render a Mermaid diagram to PNG via mmdc and embed in the document."""
    if not definition.strip():
        print("Warning: Empty Mermaid definition, skipping")
        return

    tmp_input = None
    tmp_output = None
    try:
        # Write definition to temp .mmd file
        tmp_input = tempfile.NamedTemporaryFile(suffix='.mmd', mode='w', delete=False, encoding='utf-8')
        tmp_input.write(definition)
        tmp_input.close()

        # Create temp output path
        tmp_output_path = tmp_input.name.replace('.mmd', '.png')

        # Calculate pixel width at 150 DPI
        pixel_width = int(width * 150)

        # Run mmdc to render Mermaid to PNG
        cmd = [
            'npx', '-y', '@mermaid-js/mermaid-cli',
            '-i', tmp_input.name,
            '-o', tmp_output_path,
            '-b', 'white',
            '-w', str(pixel_width),
            '-s', '2',
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode == 0 and Path(tmp_output_path).exists():
            add_image(doc, tmp_output_path, width, caption)
        else:
            error_msg = result.stderr.strip() if result.stderr else "Unknown error"
            print(f"Warning: Mermaid rendering failed: {error_msg}")
            # Fallback: add raw definition as code-style text
            para = doc.add_paragraph()
            run = para.add_run("[Mermaid diagram could not be rendered]")
            run.bold = True
            run.font.size = Pt(10)
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(definition)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(8)
            code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if caption:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.italic = True
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = DONYATI_PURPLE

    except subprocess.TimeoutExpired:
        print("Warning: Mermaid rendering timed out after 30 seconds")
        para = doc.add_paragraph()
        run = para.add_run("[Mermaid diagram rendering timed out]")
        run.bold = True
        run.font.size = Pt(10)
    except FileNotFoundError:
        print("Warning: npx not found - install Node.js to render Mermaid diagrams")
        para = doc.add_paragraph()
        run = para.add_run("[Mermaid rendering requires Node.js/npx]")
        run.bold = True
        run.font.size = Pt(10)
    finally:
        # Clean up temp files
        if tmp_input:
            try:
                Path(tmp_input.name).unlink(missing_ok=True)
            except OSError:
                pass
            try:
                png_path = tmp_input.name.replace('.mmd', '.png')
                Path(png_path).unlink(missing_ok=True)
            except OSError:
                pass


def process_content_item(doc: Document, item: Dict[str, Any]):
    """Process a single content item."""
    item_type = item.get("type", "paragraph")

    if item_type == "paragraph":
        add_paragraph(doc, item.get("text", ""))

    elif item_type == "bullets":
        add_bullets(doc, item.get("items", []))

    elif item_type == "numbered":
        add_numbered_list(doc, item.get("items", []))

    elif item_type == "table":
        add_table(doc, item.get("headers", []), item.get("rows", []))

    elif item_type == "heading":
        add_heading(doc, item.get("text", ""), item.get("level", 2))

    elif item_type == "callout":
        add_callout_box(
            doc,
            item.get("title", ""),
            item.get("text", ""),
            item.get("style", "info")
        )

    elif item_type == "highlight_box":
        add_highlight_box(doc, item.get("text", ""), item.get("style", "primary"))

    elif item_type == "metrics":
        add_metric_row(doc, item.get("items", []))

    elif item_type == "image":
        add_image(doc, item.get("path", ""), item.get("width", 6.5), item.get("caption"))

    elif item_type == "mermaid":
        add_mermaid_diagram(
            doc,
            item.get("definition", ""),
            item.get("width", 6.5),
            item.get("caption"),
        )


def create_document(content_path: str, output_path: str, template_override: Optional[str] = None):
    """Create a Word document from content specification."""

    with open(content_path, 'r', encoding='utf-8') as f:
        content = json.load(f)

    # Determine template
    if template_override == "cover":
        use_cover = True
    elif template_override == "standard":
        use_cover = False
    else:
        use_cover = should_use_cover_template(content)

    template_path = COVER_TEMPLATE if use_cover else STANDARD_TEMPLATE

    if not template_path.exists():
        print(f"Error: Template not found: {template_path}")
        sys.exit(1)

    print(f"Using template: {'Cover Page' if use_cover else 'Standard'}")
    print(f"Template file: {template_path.name}")

    # Open template
    doc = Document(str(template_path))

    # Replace cover page placeholders BEFORE removing content
    title = content.get("title", "")
    client = content.get("client", "")
    if use_cover:
        replace_cover_page_placeholders(doc, title, client)

    # Update header/footer
    # header_text replaces "Enter Report Name Here." in header
    # date replaces "Click or tap to enter date." in footer
    # footer_text replaces center copyright text if provided
    update_header(doc, content.get("header_text", content.get("title")))
    update_footer(doc, content.get("footer_text"), content.get("date"))

    # Clear template sample content (keep styles and cover page elements)
    # Standard template has style guide content that must be removed
    # Cover template has cover page elements that must be preserved

    # Known style guide / sample content patterns to remove
    sample_content_patterns = [
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "Arial", "Pangram", "Font Color", "18pt", "16pt", "14pt", "12pt",
        "#4A4778", "#12002A", "Donyati Black",
        "Bullet 1", "Bullet 2", "Bullet 3", "Bullet 4",
        "Sample Table", "Normal type", "type (",
        "Text 12pt"
    ]

    paragraphs_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Skip empty paragraphs (keep for spacing)
        if not text:
            continue

        # Keep cover page elements (title, client after replacement)
        if text in [title, client, f"Prepared For: {client}", "Prepared For:"]:
            continue

        # Keep copyright/branding
        if "©" in text or "All Rights Reserved" in text:
            continue

        # Check if this looks like sample/style guide content
        is_sample_content = False
        for pattern in sample_content_patterns:
            if pattern in text:
                is_sample_content = True
                break

        # For cover template, only remove obvious sample content
        # For standard template, remove all non-essential content
        if use_cover:
            if is_sample_content:
                paragraphs_to_remove.append(para)
        else:
            # Standard template: remove all sample content
            paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

    # Remove sample tables
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

    # Remove leading empty paragraphs (leftover space from removed content)
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Add title as Heading 1 (in document body, not cover page)
    if title:
        add_heading(doc, title, 1)

    # Add subtitle if provided
    subtitle = content.get("subtitle")
    if subtitle:
        p = doc.add_paragraph(subtitle, style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add date if provided
    date = content.get("date")
    if date:
        p = doc.add_paragraph(date, style="Normal")
        doc.add_paragraph()

    # Process sections
    sections = content.get("sections", [])
    for section in sections:
        heading = section.get("heading")
        level = section.get("level", 2)
        if heading:
            add_heading(doc, heading, level)

        section_content = section.get("content", [])
        for item in section_content:
            process_content_item(doc, item)

    # Save document
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))

    print(f"\nCreated: {output_path}")
    print(f"Sections: {len(sections)}")
    if client:
        print(f"Client: {client}")


def main():
    parser = argparse.ArgumentParser(
        description="Create Donyati-branded Word document from JSON content spec.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python create_document.py content.json report.docx
  python create_document.py content.json memo.docx --template standard
  python create_document.py proposal.json proposal.docx --template cover
        """
    )
    parser.add_argument("content", help="Path to content JSON file")
    parser.add_argument("output", help="Path for output DOCX file")
    parser.add_argument(
        "--template", "-t",
        choices=["standard", "cover"],
        help="Override template selection (default: auto-detect)"
    )

    args = parser.parse_args()

    content_path = Path(args.content)
    if not content_path.exists():
        print(f"Error: Content file not found: {args.content}")
        sys.exit(1)

    try:
        create_document(str(content_path), args.output, args.template)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
